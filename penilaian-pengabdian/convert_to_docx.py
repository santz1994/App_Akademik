import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color, qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

def add_table(doc, header_line, data_lines):
    headers = [h.strip() for h in header_line.strip().strip('|').split('|')]
    rows = []
    for line in data_lines:
        if line.strip().startswith('|---') or not line.strip().startswith('|'):
            continue
        cols = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cols)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1e293b')

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            if ci < len(headers):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(9)

    doc.add_paragraph()

with open(r'd:\Project\Blanko Penilaian Pengabdian\penilaian-pengabdian\TasksDone.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip()

    if line.startswith('# ') and not line.startswith('## '):
        doc.add_heading(line[2:], level=0)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('> '):
        p = doc.add_paragraph(line[2:])
        if p.runs:
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    elif line.startswith('| ') and i + 1 < len(lines) and lines[i+1].strip().startswith('|---'):
        table_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        add_table(doc, table_lines[0], table_lines[1:])
        continue
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.strip():
        doc.add_paragraph(line)

    i += 1

output = r'd:\Project\Blanko Penilaian Pengabdian\penilaian-pengabdian\TasksDone.docx'
doc.save(output)
print(f'Berhasil: {output}')
